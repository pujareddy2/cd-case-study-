import docx
from docx.shared import Inches

def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        # Simplistic replacement, might lose some intra-paragraph formatting if run boundaries are split
        # but safe for names. We do it at the run level if possible, else just paragraph.text.
        # Actually, replacing paragraph.text clears all runs and formatting.
        # We'll just do it carefully.
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)

def rewrite_document():
    doc = docx.Document("temp_converted.docx")
    
    # 1. Names to replace based on the extracted text
    # We will look for "Pedaprolu", "Jessica", "Vignitha" etc and replace the entire paragraph or run.
    new_names = [
        "Midde Puja \u2013 160623733160",
        "Mudakala Ujasvi \u2013 160623733162",
        "Ennala Sreenidhi \u2013 160623733144",
        "Shaik Afreen Tarunam \u2013 160623733176",
        "M. Anjali \u2013 160623733154",
        "R. Anusha \u2013 160623733167"
    ]
    new_names_str = "\n".join(new_names)
    
    # Flags and state
    replaced_front_names = False
    replaced_cert_names = False
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        
        # Replace student names
        if "Pedaprolu" in text or "160624733315" in text:
            if not replaced_front_names:
                p.text = new_names_str
                replaced_front_names = True
            else:
                p.text = "" # Clear duplicates
                
        if "Jessica" in text or "Vignitha" in text or "160623733146" in text or "160623733150" in text:
            if not replaced_cert_names:
                p.text = new_names_str
                replaced_cert_names = True
            else:
                p.text = ""

        # Content Rewriting based on keywords
        if "Stack allocation" in text or "Activation Record" in text and len(text) > 50:
            p.text = "An activation record, or stack frame, is dynamically instantiated onto the executionStack whenever a procedure is invoked. This data structure encapsulates the essential state required for the function's execution, including local variables, formal parameters, return addresses, and control links (dynamic links) referencing the caller's frame. The runtime support environment orchestrates the pushing and popping of these records, meticulously aligning with the Last-In-First-Out (LIFO) semantics of the call stack."
            
        elif "Heap allocation" in text or "Garbage Collection" in text and len(text) > 50:
            p.text = "Conversely, the dynamicMemory (heap) area services dynamic memory allocation requests. Unlike the executionStack, objects residing here lack deterministic lifespans bound by scope exits. To prevent memory exhaustion, an automated memory reclamation process, known as Garbage Collection, is deployed. Algorithms such as Mark-Sweep trace reachable objects and purge inaccessible entities, seamlessly executed by the runtime's background daemons."
            
        elif "callStack" in text:
            p.text = p.text.replace("callStack", "executionStack")
            
        elif "heap" in text.lower() and "memory" not in text.lower() and len(text) < 30:
            pass # Keep headings
        elif "heap" in text and len(text) > 20:
            p.text = p.text.replace("heap", "dynamicMemory")
            
        elif "updateStack()" in text:
            p.text = p.text.replace("updateStack()", "refreshStackView()")
            
    # Add new output screenshots
    doc.add_heading('OUTPUT SCREENSHOTS', level=1)
    
    doc.add_heading('Scenario 1: Function Calls', level=2)
    doc.add_paragraph('Visualizing the executionStack during consecutive nested calls (main, display, calculate):')
    doc.add_picture('scenario1.png', width=Inches(6.0))
    
    doc.add_heading('Scenario 2: Heap Allocation', level=2)
    doc.add_paragraph('Instantiating dynamic entities in the dynamicMemory (Customer, Student, Employee):')
    doc.add_picture('scenario2.png', width=Inches(6.0))
    
    doc.add_heading('Scenario 3: Garbage Collection', level=2)
    doc.add_paragraph('Simulating memory reclamation where an unreachable object is removed from dynamicMemory:')
    doc.add_picture('scenario3.png', width=Inches(6.0))
    
    # Add References
    doc.add_heading('REFERENCES', level=1)
    references = [
        "1. Aho, A. V., Lam, M. S., Sethi, R., & Ullman, J. D. Compilers: Principles, Techniques, and Tools (Dragon Book).",
        "2. Silberschatz, A., Galvin, P. B., & Gagne, G. Operating System Concepts.",
        "3. Appel, A. W. Modern Compiler Implementation.",
        "4. Oracle Java Documentation.",
        "5. NPTEL Course Materials on Runtime Environment.",
        "6. GeeksforGeeks articles on Memory Management.",
        "7. IEEE Research Papers on Memory Allocation and Garbage Collection Optimization."
    ]
    for ref in references:
        doc.add_paragraph(ref)
        
    doc.save("Final_Report_CD_Case_Study_3.docx")

if __name__ == "__main__":
    rewrite_document()
    print("Report written successfully.")
