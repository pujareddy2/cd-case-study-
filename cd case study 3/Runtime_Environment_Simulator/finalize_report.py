import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

def finalize_report():
    print("Loading original converted document...")
    doc = docx.Document("temp_converted.docx")

    # The user wants to replace student names perfectly.
    # The original PDF had Pedaprolu, Jessica, Vignitha etc.
    new_names = [
        "Midde Puja \u2013 160623733160",
        "Mudakala Ujasvi \u2013 160623733162",
        "Ennala Sreenidhi \u2013 160623733144",
        "Shaik Afreen Tarunam \u2013 160623733176",
        "M. Anjali \u2013 160623733154",
        "R. Anusha \u2013 160623733167"
    ]
    
    # We will search the paragraphs for old names and clear them,
    # except the first one we find where we will inject the new list.
    old_name_identifiers = ["Pedaprolu", "160624733315", "Jessica", "Vignitha", "160623733146", "160623733150"]
    
    front_replaced = False
    cert_replaced = False

    for p in doc.paragraphs:
        # Check if paragraph contains an old name
        if any(name in p.text for name in old_name_identifiers):
            # Is it the front page (first half of doc) or certificate page?
            # We'll just replace the first hit with all 6 names, and clear subsequent hits.
            if not front_replaced:
                p.text = ""
                for n in new_names:
                    run = p.add_run(n + "\n")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                front_replaced = True
            elif not cert_replaced and "Jessica" in p.text:
                p.text = ""
                for n in new_names:
                    run = p.add_run(n + "\n")
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                cert_replaced = True
            else:
                p.text = "" # Clear duplicate old names

        # Set tab indent for actual content paragraphs
        if len(p.text) > 100:
            p.paragraph_format.first_line_indent = Inches(0.5)

    # ---------------------------------------------------------
    # Expand Content
    # ---------------------------------------------------------
    print("Expanding content...")
    doc.add_heading('Advanced Runtime Memory Concepts', level=1)
    
    doc.add_heading('Stack vs Heap Allocation Advantages & Disadvantages', level=2)
    p = doc.add_paragraph("Stack Allocation Advantages: Highly deterministic, extremely fast O(1) allocation/deallocation via simple pointer arithmetic. Automatically managed by the compiler.\n")
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    p = doc.add_paragraph("Stack Allocation Disadvantages: Size is strictly limited and must be known at compile time (or limited by OS limits, causing Stack Overflow). Variables cannot persist beyond the function's scope.\n")
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    p = doc.add_paragraph("Heap Allocation Advantages: Allows dynamic sizing at runtime. Variables can outlive the function that created them, enabling complex data structures like linked lists and trees.\n")
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    p = doc.add_paragraph("Heap Allocation Disadvantages: Much slower due to complex search algorithms for free blocks. Requires explicit deallocation or Garbage Collection, both of which introduce performance overhead. Susceptible to memory leaks and fragmentation.\n")
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_heading('Real-Time Applications', level=2)
    p = doc.add_paragraph("The concepts simulated in this project are deployed in real-world systems. For example, the Java Virtual Machine (JVM) utilizes a highly optimized Garbage Collector (like G1GC or ZGC) running in background threads to manage the dynamicMemory heap. Similarly, the V8 JavaScript Engine relies on hidden classes and mark-and-sweep algorithms to handle rapid dynamic object allocation in web browsers. Real-time systems often use deterministic Garbage Collectors to avoid unpredictable 'stop-the-world' pauses.")
    p.paragraph_format.first_line_indent = Inches(0.5)

    # ---------------------------------------------------------
    # Output Screenshots
    # ---------------------------------------------------------
    print("Adding screenshots...")
    doc.add_heading('OUTPUT SCREENSHOTS', level=1)
    
    try:
        doc.add_heading('Scenario 1: Function Calls', level=2)
        doc.add_paragraph('Visualizing the executionStack during consecutive nested calls (main, display, calculate):')
        doc.add_picture('scenario1.png', width=Inches(6.0))
        
        doc.add_heading('Scenario 2: Heap Allocation', level=2)
        doc.add_paragraph('Instantiating dynamic entities in the dynamicMemory (Customer, Student, Employee):')
        doc.add_picture('scenario2.png', width=Inches(6.0))
        
        doc.add_heading('Scenario 3: Garbage Collection', level=2)
        doc.add_paragraph('Simulating memory reclamation where an unreachable object is removed from dynamicMemory:')
        doc.add_picture('scenario3.png', width=Inches(6.0))
    except Exception as e:
        print(f"Warning: Could not add screenshots: {e}")

    # ---------------------------------------------------------
    # Code Snippets
    # ---------------------------------------------------------
    doc.add_heading('CODE SNIPPETS', level=1)
    doc.add_paragraph("Snippet: Garbage Collection (Mark-Sweep Algorithm)")
    code1 = """
    def sweep(self) -> int:
        unreachable = []
        for obj_id, obj in self.hm.heap.items():
            if not obj.marked:
                unreachable.append(obj_id)
                
        for obj_id in unreachable:
            del self.hm.heap[obj_id]
            
        return len(unreachable)
    """
    p = doc.add_paragraph(code1)
    p.style = 'No Spacing'
    p.runs[0].font.name = 'Courier New'

    # Add References
    doc.add_heading('REFERENCES', level=1)
    references = [
        "1. Aho, A. V., Lam, M. S., Sethi, R., & Ullman, J. D. Compilers: Principles, Techniques, and Tools (Dragon Book).",
        "2. Silberschatz, A., Galvin, P. B., & Gagne, G. Operating System Concepts.",
        "3. Appel, A. W. Modern Compiler Implementation.",
        "4. Oracle Java Documentation.",
        "5. NPTEL Course Materials on Runtime Environment.",
        "6. GeeksforGeeks articles on Memory Management.",
        "7. IEEE Research Papers on Memory Allocation and Garbage Collection Optimization."
    ]
    for ref in references:
        doc.add_paragraph(ref)
        
    # Add Page Numbers to the footer
    print("Adding page numbers...")
    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'separate')

        fldChar3 = create_element('w:fldChar')
        create_attribute(fldChar3, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        add_page_number(run)

    doc.save("Final_Report_CD_Case_Study_3.docx")
    print("Final_Report_CD_Case_Study_3.docx saved successfully.")

if __name__ == "__main__":
    finalize_report()
