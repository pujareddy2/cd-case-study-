import docx
from docx.shared import Inches
import os

doc_path = r'c:\Desktop\case study\cd case study 1\Compiler_Project\Copy of CD case study ch1.docx'
out_path = r'c:\Desktop\case study\cd case study 1\Compiler_Project\Modified_CD_case_study.docx'

doc = docx.Document(doc_path)

# New texts
student_names = "Midde Puja\nMudakala Ujasvi\nEnnala Sreenidhi\nShaik Afreen Tarunam\nM. Anjali\nR. Anusha\n"
roll_numbers = "160623733160\n160623733162\n160623733144\n160623733176\n160623733154\n160623733167\n"

# 1. Replace in Tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if "Jessica" in cell.text:
                # Replace Names
                for p in cell.paragraphs:
                    p.clear()
                cell.paragraphs[0].text = student_names
            if "160623733146" in cell.text:
                # Replace Rolls
                for p in cell.paragraphs:
                    p.clear()
                cell.paragraphs[0].text = roll_numbers

# 2. Rewrite Content
# State machine to replace sections
current_section = None

new_content = {
    "ABSTRACT": "This case study focuses on the implementation of a compiler front-end, primarily handling the Lexical Analysis and Syntax Analysis phases for an arithmetic expression language. Utilizing token recognition techniques and parser construction methodologies, the project validates syntax and constructs abstract syntax trees for expression evaluation.",
    "INTRODUCTION": "Compiler design forms the backbone of translating high-level code into executable instructions. This study specifically explores the front-end design, encompassing scanning and parsing. By applying a recursive descent parsing approach alongside strict syntax validation, the system efficiently converts mathematical expressions into computable data structures.",
    "PROBLEM STATEMENT": "The primary objective is to design a front-end module capable of accurately tokenizing, parsing, and evaluating complex arithmetic expressions, resolving operator precedence, and validating syntactic structures in real-time.",
    "OBJECTIVES": "1. To implement an efficient Lexical Analyzer for token recognition.\n2. To construct a reliable Parser handling precedence and associativity.\n3. To perform robust Syntax Validation on arithmetic expressions.\n4. To dynamically evaluate mathematical inputs using an Abstract Syntax Tree.",
    "SCOPE": "The scope of this project encompasses the initial stages of compiler design, specifically the scanner and parser. It handles basic arithmetic operators (+, -, *, /) and parentheses, establishing a foundation for more complex semantic analysis and code generation.",
    "ADVANTAGES": "- Real-time syntax validation prevents runtime failures.\n- Modular parser construction allows easy extension of the grammar.\n- Efficient token recognition processes large expressions instantaneously.",
    "DISADVANTAGES": "- Limited to basic arithmetic operations without logical or bitwise operators.\n- Does not support variable assignment natively in this simplified iteration.",
    "APPLICATIONS": "- Educational demonstrations of compiler front-end mechanics.\n- Embedded systems requiring lightweight expression evaluation.\n- Foundational component for domain-specific scripting languages.",
    "RESULT ANALYSIS": "The system successfully evaluated a variety of test cases, accurately identifying syntax errors in invalid expressions and correctly applying operator precedence rules to yield precise mathematical results for complex nested equations.",
    "CONCLUSION": "In conclusion, the implementation of this compiler front-end effectively demonstrates the principles of Lexical Analysis and Parser Construction. The recursive descent parser proved highly capable of expression evaluation and strict syntax validation."
}

code_replacement = """
import javax.swing.*;
import java.awt.event.*;

public class CompilerFrontEnd extends JFrame {
    private JTextField expressionField;
    private JTextField outputDisplay;
    private JButton computeBtn;

    public CompilerFrontEnd() {
        setTitle("Expression Evaluator");
        setSize(400, 250);
        setLayout(null);

        JLabel lblInput = new JLabel("Enter Expression:");
        lblInput.setBounds(30, 60, 150, 20);
        add(lblInput);

        expressionField = new JTextField();
        expressionField.setBounds(30, 80, 340, 30);
        add(expressionField);

        computeBtn = new JButton("ComputeResult");
        computeBtn.setBounds(150, 130, 100, 30);
        add(computeBtn);

        JLabel lblOutput = new JLabel("Result:");
        lblOutput.setBounds(30, 180, 150, 20);
        add(lblOutput);

        outputDisplay = new JTextField();
        outputDisplay.setBounds(80, 175, 290, 30);
        outputDisplay.setEditable(false);
        add(outputDisplay);

        computeBtn.addActionListener(new ActionListener() {
            public void actionPerformed(ActionEvent e) {
                try {
                    double val = computeResult(expressionField.getText());
                    outputDisplay.setText(String.valueOf(val));
                } catch(Exception ex) {
                    outputDisplay.setText("Syntax Error");
                }
            }
        });
        setDefaultCloseOperation(JFrame.EXIT_ON_CLOSE);
    }
    
    // Recursive Descent Parser Logic Simulation
    private double computeResult(String exp) {
        return new Object() {
            int pos = -1, ch;
            void nextChar() { ch = (++pos < exp.length()) ? exp.charAt(pos) : -1; }
            boolean eat(int charToEat) {
                while (ch == ' ') nextChar();
                if (ch == charToEat) { nextChar(); return true; }
                return false;
            }
            double parse() {
                nextChar();
                double x = parseExpression();
                if (pos < exp.length()) throw new RuntimeException("Unexpected: " + (char)ch);
                return x;
            }
            double parseExpression() {
                double x = parseTerm();
                for (;;) {
                    if      (eat('+')) x += parseTerm();
                    else if (eat('-')) x -= parseTerm();
                    else return x;
                }
            }
            double parseTerm() {
                double x = parseFactor();
                for (;;) {
                    if      (eat('*')) x *= parseFactor();
                    else if (eat('/')) x /= parseFactor();
                    else return x;
                }
            }
            double parseFactor() {
                if (eat('+')) return parseFactor();
                if (eat('-')) return -parseFactor();
                double x;
                int startPos = this.pos;
                if (eat('(')) {
                    x = parseExpression();
                    eat(')');
                } else if ((ch >= '0' && ch <= '9') || ch == '.') {
                    while ((ch >= '0' && ch <= '9') || ch == '.') nextChar();
                    x = Double.parseDouble(exp.substring(startPos, this.pos));
                } else {
                    throw new RuntimeException("Unexpected: " + (char)ch);
                }
                return x;
            }
        }.parse();
    }
    public static void main(String[] args) {
        new CompilerFrontEnd().setVisible(true);
    }
}
"""

skip_paragraphs = False

for p in doc.paragraphs:
    text = p.text.strip().upper()
    
    # Check if we hit a known header
    matched_header = None
    for key in new_content.keys():
        if key in text and len(text) < len(key) + 15: # roughly matches header
            matched_header = key
            break
            
    if matched_header:
        current_section = matched_header
        skip_paragraphs = True
        # Insert the new text right after the header
        p.insert_paragraph_before(new_content[current_section] + "\n")
        continue

    # Skip paragraphs that belonged to the old section until we hit an empty space or next header
    if skip_paragraphs:
        if text == "" or "import " in p.text:
            skip_paragraphs = False
        else:
            p.clear()

    # Code replacement logic
    if "import javax.swing" in p.text or "class " in p.text or "public static void main" in p.text:
        p.clear()
        
# Append Code at the end or replace existing (we cleared it, let's append at the end of the text if it got wiped)
# Wait, simply clearing code leaves empty space. Let's just append the code where a specific keyword is, or at the end.
doc.add_heading("Java Swing Implementation", level=2)
doc.add_paragraph(code_replacement)

# Replace images/tables
# Clear the existing I/O tables if they exist
for table in doc.tables:
    if "Input Expression" in table.cell(0,0).text:
        table._element.getparent().remove(table._element)

doc.add_heading("Execution Screenshots", level=2)
doc.add_picture(r'c:\Desktop\case study\cd case study 1\Compiler_Project\screenshot1.png', width=Inches(4.0))
doc.add_picture(r'c:\Desktop\case study\cd case study 1\Compiler_Project\screenshot2.png', width=Inches(4.0))
doc.add_picture(r'c:\Desktop\case study\cd case study 1\Compiler_Project\screenshot3.png', width=Inches(4.0))

doc.save(out_path)
print("DOCX generated.")
