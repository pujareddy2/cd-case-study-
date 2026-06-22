import docx
from docx.shared import Inches, Pt
import os

doc_path = r'c:\Desktop\case study\cd case study 1\Compiler_Project\Copy of CD case study ch1.docx'
out_path = r'c:\Desktop\case study\cd case study 1\Compiler_Project\Modified_CD_case_study.docx'

doc = docx.Document(doc_path)

# New texts
student_names = "Midde Puja\nMudakala Ujasvi\nEnnala Sreenidhi\nShaik Afreen Tarunam\nM. Anjali\nR. Anusha\n"
roll_numbers = "160623733160\n160623733162\n160623733144\n160623733176\n160623733154\n160623733167\n"

def set_text_preserve_format(paragraph, text):
    # Capture font from first run if available
    font_name = 'Cambria'
    font_size = None
    bold = False
    if paragraph.runs:
        font_name = paragraph.runs[0].font.name or 'Cambria'
        font_size = paragraph.runs[0].font.size
        bold = paragraph.runs[0].font.bold

    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = font_name
    if font_size:
        run.font.size = font_size
    run.font.bold = bold

# 1. Replace in Tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if "Jessica" in cell.text:
                set_text_preserve_format(cell.paragraphs[0], student_names)
                for p in cell.paragraphs[1:]:
                    p.clear()
            if "160623733146" in cell.text:
                set_text_preserve_format(cell.paragraphs[0], roll_numbers)
                for p in cell.paragraphs[1:]:
                    p.clear()

new_content = {
    "ABSTRACT": "This case study focuses on the implementation of a compiler front-end, primarily handling the Lexical Analysis and Syntax Analysis phases for an arithmetic expression language. Utilizing token recognition techniques and parser construction methodologies, the project validates syntax and constructs abstract syntax trees for expression evaluation.",
    "INTRODUCTION": "Compiler design forms the backbone of translating high-level code into executable instructions. This study specifically explores the front-end design, encompassing scanning and parsing. By applying a recursive descent parsing approach alongside strict syntax validation, the system efficiently converts mathematical expressions into computable data structures.",
    "PROBLEM STATEMENT": "The primary objective is to design a front-end module capable of accurately tokenizing, parsing, and evaluating complex arithmetic expressions, resolving operator precedence, and validating syntactic structures in real-time.",
    "OBJECTIVES": "1. To implement an efficient Lexical Analyzer for token recognition.\n2. To construct a reliable Parser handling precedence and associativity.\n3. To perform robust Syntax Validation on arithmetic expressions.\n4. To dynamically evaluate mathematical inputs using an Abstract Syntax Tree.",
    "SCOPE": "The scope of this project encompasses the initial stages of compiler design, specifically the scanner and parser. It handles basic arithmetic operators (+, -, *, /) and parentheses, establishing a foundation for more complex semantic analysis and code generation.",
    "ADVANTAGES": "- Real-time syntax validation prevents runtime failures.\n- Modular parser construction allows easy extension of the grammar.\n- Efficient token recognition processes large expressions instantaneously.",
    "DISADVANTAGES": "- Limited to basic arithmetic operations without logical or bitwise operators.\n- Does not support variable assignment natively in this simplified iteration.",
    "APPLICATIONS": "- Educational demonstrations of compiler front-end mechanics.\n- Embedded systems requiring lightweight expression evaluation.\n- Foundational component for domain-specific scripting languages.",
    "RESULT ANALYSIS": "The system successfully evaluated a variety of test cases, accurately identifying syntax errors in invalid expressions and correctly applying operator precedence rules to yield precise mathematical results for complex nested equations.",
    "CONCLUSION": "In conclusion, the implementation of this compiler front-end effectively demonstrates the principles of Lexical Analysis and Parser Construction. The recursive descent parser proved highly capable of expression evaluation and strict syntax validation."
}

code_replacement = """
import javax.swing.*;
import java.awt.event.*;

public class CompilerFrontEnd extends JFrame {
    private JTextField expressionField;
    private JTextField outputDisplay;
    private JButton computeBtn;

    public CompilerFrontEnd() {
        setTitle("Expression Evaluator");
        setSize(400, 250);
        setLayout(null);

        JLabel lblInput = new JLabel("Enter Expression:");
        lblInput.setBounds(30, 60, 150, 20);
        add(lblInput);

        expressionField = new JTextField();
        expressionField.setBounds(30, 80, 340, 30);
        add(expressionField);

        computeBtn = new JButton("ComputeResult");
        computeBtn.setBounds(150, 130, 150, 30);
        add(computeBtn);

        JLabel lblOutput = new JLabel("Result:");
        lblOutput.setBounds(30, 180, 150, 20);
        add(lblOutput);

        outputDisplay = new JTextField();
        outputDisplay.setBounds(80, 175, 290, 30);
        outputDisplay.setEditable(false);
        add(outputDisplay);

        computeBtn.addActionListener(new ActionListener() {
            public void actionPerformed(ActionEvent e) {
                try {
                    double val = computeResult(expressionField.getText());
                    outputDisplay.setText(String.valueOf(val));
                } catch(Exception ex) {
                    outputDisplay.setText("Syntax Error");
                }
            }
        });
        setDefaultCloseOperation(JFrame.EXIT_ON_CLOSE);
    }
    
    // Recursive Descent Parser Logic Simulation
    private double computeResult(String exp) {
        return new Object() {
            int pos = -1, ch;
            void nextChar() { ch = (++pos < exp.length()) ? exp.charAt(pos) : -1; }
            boolean eat(int charToEat) {
                while (ch == ' ') nextChar();
                if (ch == charToEat) { nextChar(); return true; }
                return false;
            }
            double parse() {
                nextChar();
                double x = parseExpression();
                if (pos < exp.length()) throw new RuntimeException("Unexpected: " + (char)ch);
                return x;
            }
            double parseExpression() {
                double x = parseTerm();
                for (;;) {
                    if      (eat('+')) x += parseTerm();
                    else if (eat('-')) x -= parseTerm();
                    else return x;
                }
            }
            double parseTerm() {
                double x = parseFactor();
                for (;;) {
                    if      (eat('*')) x *= parseFactor();
                    else if (eat('/')) x /= parseFactor();
                    else return x;
                }
            }
            double parseFactor() {
                if (eat('+')) return parseFactor();
                if (eat('-')) return -parseFactor();
                double x;
                int startPos = this.pos;
                if (eat('(')) {
                    x = parseExpression();
                    eat(')');
                } else if ((ch >= '0' && ch <= '9') || ch == '.') {
                    while ((ch >= '0' && ch <= '9') || ch == '.') nextChar();
                    x = Double.parseDouble(exp.substring(startPos, this.pos));
                } else {
                    throw new RuntimeException("Unexpected: " + (char)ch);
                }
                return x;
            }
        }.parse();
    }
    public static void main(String[] args) {
        new CompilerFrontEnd().setVisible(true);
    }
}
"""

# State machine for replacing text sections
current_section = None
skip_paragraphs = False

for p in doc.paragraphs:
    text = p.text.strip().upper()
    
    matched_header = None
    for key in new_content.keys():
        # Match if the key is the exact header or very close
        if key == text or (key in text and len(text) < len(key) + 15):
            matched_header = key
            break
            
    if matched_header:
        current_section = matched_header
        skip_paragraphs = True
        continue

    if skip_paragraphs:
        if text == "" or "IMPORT " in text or "CLASS " in text:
            # We hit code or empty space, end section replacement
            if text == "":
                # Don't skip empty lines, just stop skipping content
                continue
            skip_paragraphs = False
        else:
            # This is the first content paragraph of the section
            set_text_preserve_format(p, new_content[current_section])
            # Reset so we don't overwrite the next paragraph with the same text
            current_section = None
            skip_paragraphs = False
            continue

    # Code replacement logic
    if "import javax.swing" in p.text or "import java." in p.text or "class " in p.text or "public static void main" in p.text:
        p.clear()

# Locate where to insert the code
# Let's find the heading "Source Code" or similar
code_inserted = False
for p in doc.paragraphs:
    if "SOURCE CODE" in p.text.upper():
        # Insert code in the next paragraph
        p.insert_paragraph_before("\n" + code_replacement + "\n")
        code_inserted = True
        break

if not code_inserted:
    # Fallback to appending
    doc.add_heading("Source Code", level=2)
    p_code = doc.add_paragraph()
    set_text_preserve_format(p_code, code_replacement)

# Replace images/tables
# Find the Input/Output table and replace its content with pictures
table_to_delete = None
for table in doc.tables:
    if len(table.rows) > 0 and len(table.columns) > 0:
        try:
            if "Input" in table.cell(0,0).text and "Output" in table.cell(0,1).text:
                table_to_delete = table
        except:
            pass

if table_to_delete:
    table_to_delete._element.getparent().remove(table_to_delete._element)

# Find the Result Analysis section to append images
for i, p in enumerate(doc.paragraphs):
    if "RESULT ANALYSIS" in p.text.upper() or "OUTPUT" in p.text.upper():
        # Add pictures here
        run = doc.paragraphs[i+1].add_run() if i+1 < len(doc.paragraphs) else doc.add_paragraph().add_run()
        run.add_picture(r'c:\Desktop\case study\cd case study 1\Compiler_Project\screenshot1.png', width=Inches(4.0))
        run.add_break()
        run.add_picture(r'c:\Desktop\case study\cd case study 1\Compiler_Project\screenshot2.png', width=Inches(4.0))
        run.add_break()
        run.add_picture(r'c:\Desktop\case study\cd case study 1\Compiler_Project\screenshot3.png', width=Inches(4.0))
        break

doc.save(out_path)
print("Perfect DOCX generated.")
