import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_BREAK

doc_path = r'c:\Desktop\case study\cd case study 1\Compiler_Project\Copy of CD case study ch1.docx'
out_path = r'c:\Desktop\case study\cd case study 1\Compiler_Project\Modified_CD_case_study_v2.docx'

doc = docx.Document(doc_path)

replacements = {
    "Jessica": "Midde Puja",
    "K.Vignitha": "Mudakala Ujasvi",
    "Shaik Karishma": "Ennala Sreenidhi",
    "T Akshitha": "Shaik Afreen Tarunam",
    "T. Akshitha": "Shaik Afreen Tarunam",
    "V. Lasya": "M. Anjali",
    "Pedaprolu S S L Katyayani": "R. Anusha",
    "160623733146": "160623733160",
    "160623733150": "160623733162",
    "160623733177": "160623733144",
    "160623733185": "160623733176",
    "160623733187": "160623733154",
    "160624733315": "160623733167",
    "160623733315": "160623733167"
}

def replace_text_in_runs(paragraph):
    for run in paragraph.runs:
        for old, new in replacements.items():
            if old in run.text:
                run.text = run.text.replace(old, new)

# 1. Replace Student Names and Roll Numbers exactly in-place
for p in doc.paragraphs:
    replace_text_in_runs(p)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_text_in_runs(p)

# 2. Add Tab spaces to all paragraph starts where the user rewrote text.
# We will do a carefully targeted rewrite of the specific sections.
new_content = {
    "ABSTRACT": "\tThis case study focuses on the implementation of a compiler front-end, primarily handling the Lexical Analysis and Syntax Analysis phases for an arithmetic expression language. Utilizing token recognition techniques and parser construction methodologies, the project validates syntax and constructs abstract syntax trees for expression evaluation.",
    "INTRODUCTION": "\tCompiler design forms the backbone of translating high-level code into executable instructions. This study specifically explores the front-end design, encompassing scanning and parsing. By applying a recursive descent parsing approach alongside strict syntax validation, the system efficiently converts mathematical expressions into computable data structures.",
    "1.1 About Problem statement:": "\tThe primary objective is to design a front-end module capable of accurately tokenizing, parsing, and evaluating complex arithmetic expressions, resolving operator precedence, and validating syntactic structures in real-time.",
    "1.2 Objectives of the Problem statement": "\t1. To implement an efficient Lexical Analyzer for token recognition.\n\t2. To construct a reliable Parser handling precedence and associativity.\n\t3. To perform robust Syntax Validation on arithmetic expressions.\n\t4. To dynamically evaluate mathematical inputs using an Abstract Syntax Tree.",
    "1.3 Scope of the Problem statement": "\tThe scope of this project encompasses the initial stages of compiler design, specifically the scanner and parser. It handles basic arithmetic operators (+, -, *, /) and parentheses, establishing a foundation for more complex semantic analysis and code generation.",
    "1.4 Advantages": "\t- Real-time syntax validation prevents runtime failures.\n\t- Modular parser construction allows easy extension of the grammar.\n\t- Efficient token recognition processes large expressions instantaneously.",
    "1.5 Disadvantages": "\t- Limited to basic arithmetic operations without logical or bitwise operators.\n\t- Does not support variable assignment natively in this simplified iteration.",
    "1.6 Applications": "\t- Educational demonstrations of compiler front-end mechanics.\n\t- Embedded systems requiring lightweight expression evaluation.\n\t- Foundational component for domain-specific scripting languages.",
    "3. RESULT ANALYSIS": "\tThe system successfully evaluated a variety of test cases, accurately identifying syntax errors in invalid expressions and correctly applying operator precedence rules to yield precise mathematical results for complex nested equations.",
    "Conclusion": "\tIn conclusion, the implementation of this compiler front-end effectively demonstrates the principles of Lexical Analysis and Parser Construction. The recursive descent parser proved highly capable of expression evaluation and strict syntax validation."
}

def set_run_with_format(paragraph, text):
    font_name = 'Cambria'
    font_size = None
    if paragraph.runs:
        font_name = paragraph.runs[0].font.name or 'Cambria'
        font_size = paragraph.runs[0].font.size

    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = font_name
    if font_size:
        run.font.size = font_size

skip_paragraphs = False
current_section = None

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    upper_text = text.upper()
    
    # Check exact headers
    matched_header = None
    for key in new_content.keys():
        if text.startswith(key) or upper_text.startswith(key.upper()):
            matched_header = key
            break
            
    if matched_header:
        current_section = matched_header
        skip_paragraphs = True
        
        # Ensure Introduction starts on a new page
        if matched_header.upper() == "INTRODUCTION":
            # Add page break before this paragraph
            p.insert_paragraph_before("").add_run().add_break(WD_BREAK.PAGE)
            
        continue

    if skip_paragraphs:
        if text == "" or upper_text.startswith("1.") or upper_text.startswith("2.") or upper_text.startswith("3.") or "OUTPUT" in upper_text or "REFERENCES" in upper_text or "CODE" in upper_text:
            if text != "":
                skip_paragraphs = False
        else:
            set_run_with_format(p, new_content[current_section])
            current_section = None
            skip_paragraphs = False
            continue

# 3. Handle Java Code
java_code = """import javax.swing.*;
import java.awt.event.*;

public class CompilerFrontEnd extends JFrame {
    private JTextField expressionField;
    private JTextField outputDisplay;
    private JButton computeBtn;

    public CompilerFrontEnd() {
        setTitle("Expression Evaluator");
        setSize(400, 250);
        setLayout(null);

        JLabel lblInput = new JLabel("Enter Expression:");
        lblInput.setBounds(30, 60, 150, 20);
        add(lblInput);

        expressionField = new JTextField();
        expressionField.setBounds(30, 80, 340, 30);
        add(expressionField);

        computeBtn = new JButton("ComputeResult");
        computeBtn.setBounds(150, 130, 150, 30);
        add(computeBtn);

        JLabel lblOutput = new JLabel("Result:");
        lblOutput.setBounds(30, 180, 150, 20);
        add(lblOutput);

        outputDisplay = new JTextField();
        outputDisplay.setBounds(80, 175, 290, 30);
        outputDisplay.setEditable(false);
        add(outputDisplay);

        computeBtn.addActionListener(e -> {
            try {
                double val = computeResult(expressionField.getText());
                outputDisplay.setText(String.valueOf(val));
            } catch(Exception ex) {
                outputDisplay.setText("Syntax Error");
            }
        });
        setDefaultCloseOperation(JFrame.EXIT_ON_CLOSE);
    }
    
    private double computeResult(String exp) {
        return new Object() {
            int pos = -1, ch;
            void nextChar() { ch = (++pos < exp.length()) ? exp.charAt(pos) : -1; }
            boolean eat(int charToEat) {
                while (ch == ' ') nextChar();
                if (ch == charToEat) { nextChar(); return true; }
                return false;
            }
            double parse() {
                nextChar();
                return parseExpression();
            }
            double parseExpression() {
                double x = parseTerm();
                for (;;) {
                    if      (eat('+')) x += parseTerm();
                    else if (eat('-')) x -= parseTerm();
                    else return x;
                }
            }
            double parseTerm() {
                double x = parseFactor();
                for (;;) {
                    if      (eat('*')) x *= parseFactor();
                    else if (eat('/')) x /= parseFactor();
                    else return x;
                }
            }
            double parseFactor() {
                if (eat('+')) return parseFactor();
                if (eat('-')) return -parseFactor();
                double x;
                int startPos = this.pos;
                if (eat('(')) {
                    x = parseExpression();
                    eat(')');
                } else if ((ch >= '0' && ch <= '9') || ch == '.') {
                    while ((ch >= '0' && ch <= '9') || ch == '.') nextChar();
                    x = Double.parseDouble(exp.substring(startPos, this.pos));
                } else {
                    throw new RuntimeException("Unexpected: " + (char)ch);
                }
                return x;
            }
        }.parse();
    }
    public static void main(String[] args) {
        new CompilerFrontEnd().setVisible(true);
    }
}
"""

in_code_block = False
code_replaced = False
for p in doc.paragraphs:
    text = p.text.strip()
    if text == "JAVA" or "import javax.swing" in text:
        in_code_block = True
        
    if in_code_block:
        if text == "Output:" or text.upper() == "OUTPUT:":
            in_code_block = False
        else:
            if not code_replaced:
                set_run_with_format(p, java_code)
                code_replaced = True
            else:
                p.clear()

# 4. Handle Output Images
# Find "Output:" paragraph
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().upper() == "OUTPUT:":
        # Clear anything right after it if it was from previous messed up script
        run = p.add_run()
        run.add_break()
        run.add_picture(r'c:\Desktop\case study\cd case study 1\Compiler_Project\screenshot1.png', width=Inches(4.0))
        run.add_break()
        run.add_picture(r'c:\Desktop\case study\cd case study 1\Compiler_Project\screenshot2.png', width=Inches(4.0))
        run.add_break()
        run.add_picture(r'c:\Desktop\case study\cd case study 1\Compiler_Project\screenshot3.png', width=Inches(4.0))
        break

# Clear the input/output tables if they exist near the end
for table in doc.tables:
    if len(table.rows) > 0 and len(table.columns) > 0:
        try:
            if "Input" in table.cell(0,0).text and "Output" in table.cell(0,1).text:
                table._element.getparent().remove(table._element)
        except:
            pass

# Remove abstract extra spacing
# Let's find "ABSTRACT" and remove empty paragraphs before it, except maybe 1 or 2.
abstract_idx = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().upper() == "ABSTRACT":
        abstract_idx = i
        break

if abstract_idx > 0:
    # Remove some empty paragraphs right above abstract
    idx = abstract_idx - 1
    while idx >= 0 and doc.paragraphs[idx].text.strip() == "":
        p_element = doc.paragraphs[idx]._element
        p_element.getparent().remove(p_element)
        idx -= 1

doc.save(out_path)
print("Perfect DOCX V2 generated.")
