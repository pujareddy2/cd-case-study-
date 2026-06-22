import docx
from docx.shared import Inches

def replace_text(doc, old_text, new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if old_text in inline[i].text:
                    inline[i].text = inline[i].text.replace(old_text, new_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old_text in p.text:
                        inline = p.runs
                        for i in range(len(inline)):
                            if old_text in inline[i].text:
                                inline[i].text = inline[i].text.replace(old_text, new_text)

def rewrite_section(doc, heading, new_content):
    found = False
    for i, p in enumerate(doc.paragraphs):
        if heading.lower() in p.text.lower() and len(p.text) < len(heading) + 10:
            # We found the heading, now replace the next paragraph(s) until next heading or empty line
            j = i + 1
            while j < len(doc.paragraphs) and doc.paragraphs[j].text.strip() != "":
                # clear text
                for r in doc.paragraphs[j].runs:
                    r.text = ""
                j += 1
            # insert new content at i+1
            if i + 1 < len(doc.paragraphs):
                doc.paragraphs[i+1].text = new_content
            found = True
            break

def main():
    doc = docx.Document("Converted_Case_Study.docx")
    
    replacements = {
        "Jessica": "Midde Puja",
        "160623733146": "160623733160",
        "K.Vignitha": "Mudakala Ujasvi",
        "160623733150": "160623733162",
        "Shaik Karishma": "Ennala Sreenidhi",
        "160623733177": "160623733144",
        "T. Akshitha": "Shaik Afreen Tarunam",
        "T Akshitha": "Shaik Afreen Tarunam",
        "160623733185": "160623733176",
        "V. Lasya": "M. Anjali",
        "160623733187": "160623733154",
        "Pedaprolu S S L Katyayani": "R. Anusha",
        "Pedaprolu S S L Kat yayani": "R. Anusha",
        "160624733315": "160623733167",
        "160623733315": "160623733167"
    }
    
    # Due to docx splitting words across runs, a simple run replacement might miss some.
    # So we replace the whole paragraph text while trying to keep the first run's style.
    for old_text, new_text in replacements.items():
        for p in doc.paragraphs:
            if old_text in p.text:
                p.text = p.text.replace(old_text, new_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if old_text in p.text:
                            p.text = p.text.replace(old_text, new_text)

    # Rewrite sections
    rewrite_section(doc, "Abstract", "This project explores intermediate code generation (TAC) and Basic Block formation for control flow structures like if-else and while loops. By constructing a Control Flow Graph (CFG), we provide a strong foundation for compiler flow analysis and optimizations.")
    rewrite_section(doc, "Introduction", "Intermediate Representation (IR) separates the front end and back end of a compiler. This project focuses on translating parsed syntax into Three Address Code (TAC), which facilitates machine-independent optimizations. The intermediate code is further partitioned into Basic Blocks to model the program's control flow via a CFG.")
    rewrite_section(doc, "Objectives", "1. To implement TAC generation for conditional and looping constructs.\n2. To partition the TAC into single-entry, single-exit Basic Blocks.\n3. To construct a Control Flow Graph mapping the execution paths.")
    rewrite_section(doc, "Scope", "The project covers parsing if-else and while statements, generating their corresponding TAC, and forming a CFG representation. It serves as a foundational step for applying compiler optimization techniques.")
    rewrite_section(doc, "Advantages", "1. Optimized basic blocks simplify data-flow analysis.\n2. Clear control flow modeling via CFG enables advanced optimizations like dead-code elimination.\n3. Machine-independent TAC allows easy retargeting.")
    rewrite_section(doc, "Disadvantages", "1. Single pass compilation limits the scope of global optimizations.\n2. In-memory graph structures can consume significant resources for extremely large source files.")
    rewrite_section(doc, "Applications", "1. Used in modern compiler development.\n2. Static code analyzers utilize CFGs to detect unreachable code and logical errors.")
    rewrite_section(doc, "Result Analysis", "The implemented Java application successfully translates input structures to TAC, identifies leaders, partitions the instructions into basic blocks, and plots the CFG. Output screenshots confirm accurate flow extraction.")
    rewrite_section(doc, "Conclusion", "Control Flow Graphs and Three Address Code provide strong compiler optimization foundations. By effectively modularizing the execution flow into basic blocks, we achieve a robust intermediate representation essential for any modern compiler.")
    
    # Now for replacing images: we can't easily find exactly which image is which, 
    # but we can append the new screenshots at the end of the document, or find a keyword "Output" and insert there.
    out_found = False
    for p in doc.paragraphs:
        if "output" in p.text.lower() and "screenshot" in p.text.lower():
            out_found = True
            break
            
    if out_found:
        doc.add_paragraph("NEW OUTPUT SCREENSHOTS:")
        doc.add_picture("tac_output.png", width=Inches(5.0))
        doc.add_picture("bb_output.png", width=Inches(5.0))
        doc.add_picture("cfg_output.png", width=Inches(5.0))
    else:
        doc.add_page_break()
        doc.add_heading("OUTPUT SCREENSHOTS", level=1)
        doc.add_picture("tac_output.png", width=Inches(5.0))
        doc.add_picture("bb_output.png", width=Inches(5.0))
        doc.add_picture("cfg_output.png", width=Inches(5.0))
        
    doc.save("Final_Report_Case_Study_2.docx")
    print("Modifications completed. Saved to Final_Report_Case_Study_2.docx")

if __name__ == "__main__":
    main()
