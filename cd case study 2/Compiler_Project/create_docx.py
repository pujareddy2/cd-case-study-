import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(ns.qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def main():
    doc = docx.Document()

    # Apply global styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    def add_title(text, level=1, align=WD_ALIGN_PARAGRAPH.CENTER):
        p = doc.add_heading(text, level=level)
        p.alignment = align
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_para(text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        if bold:
            run.bold = True
        return p

    # PAGE 1: Title Page
    add_title("Intermediate Code Generation + Basic Block Formation + Control Flow Graph for If-Else and While : A Case Study\n", level=1)
    
    add_para("Case study Report submitted in partial fulfilment of the requirements for the award of the Degree of B.E in Computer Science and Engineering\n", align=WD_ALIGN_PARAGRAPH.CENTER)
    
    add_para("By\n", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para("Midde Puja - 160623733160", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para("Mudakala Ujasvi - 160623733162", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para("Ennala Sreenidhi - 160623733144", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para("Shaik Afreen Tarunam - 160623733176", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para("M. Anjali - 160623733154", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para("R. Anusha - 160623733167\n", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_para("Under the Guidance of", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para("Mrs. Gnana Prasuna\nAssistant Professor, Department of Computer Science & Engineering\n", align=WD_ALIGN_PARAGRAPH.CENTER)

    # ADD LOGO
    try:
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture("logo.png", width=Inches(3.0))
    except:
        pass

    add_para("\nDepartment of Computer Science and Engineering", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para("Stanley College of Engineering & Technology for Women(Autonomous)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para("Chapel Road, Abids, Hyderabad - 500001", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para("(Affiliated to Osmania University, Hyderabad, Approved by AICTE, Accredited by NBA & NAAC with A Grade)\n", align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # PAGE 2: Certificate
    add_para("Stanley College of Engineering & Technology for Women (Autonomous)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_para("Chapel Road, Abids, Hyderabad - 500001", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_title("CERTIFICATE", level=2)
    
    cert_text = "This is to certify that the case study report entitled 'Intermediate Code Generation + Basic Block Formation + Control Flow Graph for If-Else and While' is being submitted by:\n"
    add_para(cert_text)
    
    add_para("Midde Puja - 160623733160\nMudakala Ujasvi - 160623733162\nEnnala Sreenidhi - 160623733144\nShaik Afreen Tarunam - 160623733176\nM. Anjali - 160623733154\nR. Anusha - 160623733167\n", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_para("In partial fulfilment for the award of the Degree of Bachelor of Engineering in Computer Science & Engineering to the Osmania University, Hyderabad is a record of bonafide work carried out under my guidance and supervision. The results embodied in this Problem statement report have not been submitted to any other University or Institute for the award of any Degree or Diploma.\n")

    add_para("Guide: Mrs. Gnana Prasuna, Professor, Dept. of CSE", bold=True)
    add_para("Head of the Department: Dr. R Madana Mohana, Professor & Head, Dept. of CSE\n", bold=True)

    doc.add_page_break()

    # PAGE 3: Vision and Mission
    add_title("STANLEY COLLEGE OF ENGINEERING AND TECHNOLOGY FOR WOMEN (Autonomous)", level=2)
    add_para("Department of Computer Science & Engineering\n", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_title("Institute Vision:", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("Empowering girl students through professional integrated with values and character to make an Impact in the world.")
    
    add_title("Institute Mission:", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("M1: Enabling quality engineering education for girl students to make them competent and confident to succeed in professional practice and advanced learning.")
    add_para("M2: Providing state-of-art-facilities and resources towards world class education.")
    add_para("M3: Integrating qualities like humanity, social values, ethics, leadership towards their contribution to society.")

    add_title("Department Vision:", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("Empowering girl students with the contemporary knowledge in Computer Science and Engineering for their success in life.")

    add_title("Department Mission:", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("M1: To impart quality education for girl students to learn and practice various hardware and software platforms prevalent in industry.")
    add_para("M2: To achieve self-sustainability and overall development through Research and Development activities.")
    add_para("M3: To provide education for life by focusing on the inculcation of human & moral values through an honest and scientific approach.")
    add_para("M4: To groom students with good attitude, team work and personality skills.\n")

    doc.add_page_break()

    # PAGE 4: PEO, PO, PSO
    add_title("Programme Educational Objectives", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("PEO1: Our graduates shall have enhanced skills and contemporary knowledge in software and hardware technologies for professional excellence, towards successful employment, advanced learning and research.")
    add_para("PEO2: Our graduates shall have life-long learning attitude, innovation and creativity to master latest technologies, devise solutions for realistic and social issues in the society.")
    add_para("PEO3: Our graduates have good attitude and personality skills, ethical values, teamwork and leadership skill towards professionalism and ethical practices within the organization and the society.\n")

    add_title("Program Outcomes:", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("PO1: Engineering Knowledge: Apply knowledge of mathematics and science, with fundamentals of Computer Science & Engineering to be able to solve complex engineering problems related to CSE.")
    add_para("PO2: Problem Analysis: Identify, Formulate, review research literature and analyze complex engineering problems related to CSE.")
    add_para("PO3: Design/Development of solutions: Design solutions for complex engineering problems related to CSE.")
    add_para("PO4: Conduct Investigations of Complex problems: Use research-based knowledge to provide valid conclusions.")
    add_para("PO5: Modern Tool Usage: Create, Select and apply appropriate techniques, resources and modern engineering and IT tools.")
    add_para("PO6: The Engineer and Society: Apply Reasoning informed by the contextual knowledge.")
    add_para("PO7: Environment and Sustainability: Understand the impact of the CSE professional engineering solutions.")
    add_para("PO8: Ethics: Apply Ethical Principles and commit to professional ethics.")
    add_para("PO9: Individual and Team Work: Function effectively as an individual and as a member or leader in diverse teams.")
    add_para("PO10: Communication: Communicate effectively on complex engineering activities.")
    add_para("PO11: Project Management and Finance: Demonstrate knowledge and understanding of the engineering management principles.")
    add_para("PO12: Life-Long Learning: Recognize the need for and have the preparation and ability to engage in independent and life-long learning.\n")

    add_title("Program Specific Outcomes:", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("PSO1: Problem-Solving Skills: The ability to apply standard practices and strategies in software development.")
    add_para("PSO2: Design, Implement, Test and Evaluate a computer system, component or algorithm to meet desired needs.\n")

    doc.add_page_break()

    # PAGE 5: TOC
    add_title("TABLE OF CONTENTS", level=2)
    add_para("1. Abstract .................................................................................... 1")
    add_para("2. Introduction ................................................................................. 2")
    add_para("3. Objectives ................................................................................... 3")
    add_para("4. Scope ........................................................................................ 4")
    add_para("5. Compiling Steps .............................................................................. 5")
    add_para("6. Intermediate Code Generation ................................................................. 6")
    add_para("7. Basic Block Formation ........................................................................ 7")
    add_para("8. Control Flow Graph ........................................................................... 8")
    add_para("9. Advantages & Disadvantages ................................................................... 9")
    add_para("10. Applications ................................................................................ 10")
    add_para("11. Result Analysis ............................................................................. 11")
    add_para("12. Conclusion .................................................................................. 12")
    add_para("13. References .................................................................................. 13")
    
    # We create a new section here to start page numbering from 1
    new_section = doc.add_section()
    new_section.footer.is_linked_to_previous = False
    
    footer_para = new_section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("Page ")
    add_page_number(run)

    # PAGE 6: Technical Content
    add_title("1. Abstract", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("This project explores the crucial phases of modern compiler backend design, specifically focusing on intermediate code generation (TAC) and Basic Block formation for complex control flow structures like if-else and while loops. Compilers rely on these intermediate structures to break down high-level syntax into easily analyzable, machine-independent formats. By algorithmically partitioning the linear Three Address Code into single-entry, single-exit basic blocks, we systematically identify the distinct execution paths of a given program. Furthermore, by constructing a Control Flow Graph (CFG) from these blocks, we provide a mathematically rigorous foundation for flow analysis, which is vital for performing aggressive compiler optimizations such as dead-code elimination, loop unrolling, and constant propagation. This case study successfully demonstrates the automated transformation of conditional constructs into these robust graphical and linear representations, emphasizing their importance in modern software engineering.")

    doc.add_page_break()
    add_title("2. Introduction", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("Intermediate Representation (IR) is the bridge that separates the front end and back end of a modern compiler architecture. This separation is vital as it allows compilers to support multiple source languages and target machine architectures without needing an entirely distinct compiler for every combination. This project focuses intensely on translating parsed syntax into Three Address Code (TAC), a specific form of IR where every instruction consists of at most one operator and three operands. TAC facilitates machine-independent optimizations by simplifying complex mathematical expressions into discrete, easily manageable operations using temporary variables.\n\nOnce the TAC is generated, it must be analyzed for control flow. The intermediate code is partitioned into Basic Blocks—sequences of instructions with a single entry point and a single exit point. Identifying these blocks prevents the compiler from needing to analyze branching logic at every single line. Finally, these Basic Blocks are connected to form a Control Flow Graph (CFG). The CFG serves as the definitive structural map of a program's execution, where nodes represent the basic blocks and directed edges denote the conditional and unconditional jumps between them. Understanding and implementing these phases is fundamental to compiler construction and static code analysis.")

    doc.add_page_break()
    add_title("3. Objectives", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("1. To design and implement an automated Three Address Code (TAC) generator capable of handling complex conditional (if-else) and looping constructs seamlessly.")
    add_para("2. To develop an algorithm that systematically identifies leaders within the TAC and partitions the instruction stream into cohesive, single-entry, single-exit Basic Blocks.")
    add_para("3. To construct a highly accurate Control Flow Graph (CFG) that mathematically models the execution paths, jump targets, and conditional branching logic between the identified basic blocks.")
    add_para("4. To provide a clear, graphical Java Swing interface that visually demonstrates these transformation phases in real-time, bridging the gap between theoretical compiler design and practical software implementation.")

    doc.add_page_break()
    add_title("4. Scope", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("The scope of this project encompasses the back-end foundational logic of a compiler. It covers the parsing of standard programming constructs, primarily focusing on `if-else` and `while` statements, which traditionally introduce branching complexity. The project's boundaries include generating the corresponding linear Three Address Code, logically partitioning it into Basic Blocks, and outputting the final Control Flow Graph representation (nodes and edges). While it builds the necessary framework for advanced optimizations, the implementation of specific global optimizations (like data-flow analysis or register allocation) lies outside the immediate scope of this specific case study. It acts as an essential educational and functional tool for understanding the intermediate compilation pipeline.")

    doc.add_page_break()
    add_title("5. Compiling Steps", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("To execute the intermediate code generator and view the graphical interface, follow these rigorous steps:")
    add_para("1. Prerequisites: Ensure the Java Development Kit (JDK 8 or higher) is properly installed on your machine and that `javac` and `java` are configured in your system's PATH variable.")
    add_para("2. File Placement: Place the core source files `CompilerSwingApp.java` and `ScreenshotGenerator.java` into your designated project workspace directory.")
    add_para("3. Command Line Interface: Open a terminal (Linux/macOS) or command prompt/PowerShell (Windows) and navigate to the project directory.")
    add_para("4. Compilation: Compile the Java source code into bytecode using the Java compiler command: `javac CompilerSwingApp.java ScreenshotGenerator.java`")
    add_para("5. Execution (Interactive): Launch the interactive graphical user interface by executing: `java CompilerSwingApp`")
    add_para("6. Execution (Automated Screenshots): To autonomously execute the application, simulate the compiler flow on the example code, and save high-fidelity output PNGs directly to the directory, execute: `java ScreenshotGenerator`")
    add_para("7. Result Verification: The system will launch the UI, process the if-else example, construct the TAC and CFG, and save the corresponding visual outputs to the current folder.")

    doc.add_page_break()
    add_title("6. Intermediate Code Generation", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("We utilized a Java Swing-based compiler application to parse conditional inputs. The system generates Three Address Code (TAC) representing the logic explicitly using temporaries and jump labels. TAC breaks down nested operations into a sequence of simpler assignments, ensuring that the target machine can easily translate them into assembly language.")
    add_para("\nDetailed Execution Flow:\n1. The parser receives the mathematical expression and control structures from the input buffer.\n2. It traverses the Abstract Syntax Tree (AST) to break complex expressions down into atomic operations.\n3. Temporary variables (e.g., t1, t2) are dynamically assigned to hold the intermediate result of each atomic boolean or arithmetic operation.\n4. Conditional jumps (e.g., ifFalse) and unconditional jumps (goto) are emitted to replicate the control flow implicitly defined by the if-else blocks.\n5. Labels (e.g., L1, L2) are generated to serve as destination markers for the emitted jump instructions.")
    add_para("\nExample Input:\nif(a>b)\n    c=a+b;\nelse\n    c=a-b;\n")
    add_para("Generated TAC:\nt1=a>b\nifFalse t1 goto L1\nc=a+b\ngoto L2\nL1:\nc=a-b\nL2:\n")

    doc.add_page_break()
    add_title("7. Basic Block Formation", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("Once the TAC is generated, it must be modularized. The compiler successfully identifies leaders and partitions the linear instructions into fundamental Basic Blocks. A basic block is the longest sequence of instructions where the flow of control can only enter at the first instruction and exit at the last instruction.")
    add_para("\nDetailed Execution Flow:\n1. The partitioning algorithm scans the TAC line-by-line to identify specific instructions called 'leaders'.\n2. The very first TAC instruction is automatically classified as a leader.\n3. The target destination of any conditional or unconditional jump (e.g., a Label like L1) is marked as a leader.\n4. Any instruction immediately following a conditional or unconditional jump is marked as a leader.\n5. Basic blocks are then formed starting from a leader down to (but not including) the next leader.\n6. This rigorous separation ensures that if the first instruction of a basic block is executed, all subsequent instructions in that block will also execute in order without interruption.")

    doc.add_page_break()
    add_title("8. Control Flow Graph", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("CFG generation tracks the execution flow between Basic Blocks. Nodes represent the blocks, and edges map the jump conditions across the blocks, yielding a structured Adjacency List for the true and false branch evaluations. This graph is the cornerstone of advanced optimization.")
    add_para("\nDetailed Execution Flow:\n1. The system instantiates a graphical node for each mathematically identified Basic Block.\n2. It establishes directed edges between blocks based on the jump targets analyzed at the end of each block.\n3. If a block ends with an unconditional jump (`goto`), a single directed edge connects it exclusively to the target block.\n4. If a block ends with a conditional jump (`ifFalse`), two distinct edges are created: one for the 'true' execution path (fall-through) and one for the 'false' path (jumping to the label).\n5. Unreachable code is easily identified as any node without incoming edges (excluding the entry node).")
    
    add_para("\nGraph Representation:")
    try:
        doc.add_picture("cfg_diagram.png", width=Inches(5.5))
    except Exception as e:
        add_para(f"[CFG Diagram missing: Please ensure draw_cfg.py was executed to generate cfg_diagram.png]")


    doc.add_page_break()
    add_title("9. Advantages & Disadvantages", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("Advantages:", bold=True)
    add_para("1. Optimized basic blocks dramatically simplify data-flow and liveness analysis for local optimizations.")
    add_para("2. Clear control flow modeling via CFG enables advanced global optimizations like dead-code elimination, loop invariant code motion, and strength reduction.")
    add_para("3. Machine-independent TAC allows easy retargeting of the compiler to different hardware architectures without rewriting the front end.")
    add_para("4. Modular design makes debugging the compilation process significantly easier by isolating the structural transformation steps.")
    
    add_para("\nDisadvantages:", bold=True)
    add_para("1. Single pass compilation limits the scope of global optimizations, often necessitating multiple passes over the IR which increases compilation time.")
    add_para("2. Maintaining in-memory graph structures (nodes, edges, adjacency lists) can consume significant memory resources for extremely large, enterprise-level source files.")
    add_para("3. Generating temporaries for every atomic operation in TAC can initially bloat the intermediate code size before optimizations are applied.")

    doc.add_page_break()
    add_title("10. Applications", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("1. Modern Compiler Development: Every major modern compiler infrastructure (such as LLVM and GCC) relies extensively on Intermediate Representation, Basic Blocks, and CFGs to perform aggressive code optimizations.")
    add_para("2. Static Code Analyzers: Security tools and static analyzers utilize CFGs to detect unreachable code, infinite loops, uninitialized variables, and logical flaws without executing the program.")
    add_para("3. Decompilation and Reverse Engineering: Security researchers map out binary executables into CFGs to reverse engineer the logical flow of proprietary software or analyze malicious payloads in cybersecurity operations.")
    add_para("4. Software Testing: CFGs are directly used to calculate Cyclomatic Complexity and generate path-coverage test cases to ensure rigorous quality assurance.")

    doc.add_page_break()
    add_title("11. Result Analysis", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("The implemented Java Swing application successfully fulfilled all project objectives. It accurately parses conditional input structures and flawlessly translates them into Three Address Code. The algorithm correctly identifies instruction leaders, cleanly partitioning the operations into independent basic blocks. Finally, it dynamically plots the Control Flow Graph, generating a clear adjacency list that perfectly matches theoretical expectations.")
    add_para("By utilizing custom java.awt.Robot rendering techniques, the application is also capable of autonomously capturing high-fidelity execution evidence, bypassing the constraints of headless environments.")
    add_para("\nOUTPUT SCREENSHOTS:", bold=True)
    add_para("Input Program & Three Address Code Generation:")
    try:
        doc.add_picture("tac_output.png", width=Inches(6.0))
    except Exception as e:
        add_para(f"[Screenshot tac_output.png missing]")
        
    add_para("\nBasic Block Formation:")
    try:
        doc.add_picture("bb_output.png", width=Inches(6.0))
    except Exception as e:
        add_para(f"[Screenshot bb_output.png missing]")

    add_para("\nControl Flow Graph Representation:")
    try:
        doc.add_picture("cfg_output.png", width=Inches(6.0))
    except Exception as e:
        add_para(f"[Screenshot cfg_output.png missing]")

    doc.add_page_break()
    add_title("12. Conclusion", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("In conclusion, the successful translation of high-level code into Three Address Code, followed by its systematic organization into Basic Blocks and a Control Flow Graph, is a testament to the structured rigor of modern compiler design. These intermediate layers abstract away machine-specific constraints, allowing developers to focus entirely on algorithmic efficiency. Control Flow Graphs provide an indisputable mathematical model of execution flow, offering strong foundations for aggressive optimizations. By effectively modularizing the execution flow, we achieve a highly robust intermediate representation that is universally essential for any sophisticated, modern compiler architecture. This case study successfully reinforced these critical concepts through practical, graphical implementation.")

    doc.add_page_break()
    # PAGE 9: References
    add_title("13. References", level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("1. Dragon Book (Compilers: Principles, Techniques, and Tools)")
    add_para("2. Modern Compiler Implementation")
    add_para("3. Compiler Construction by Louden")
    add_para("4. NPTEL Compiler Design")
    add_para("5. GeeksforGeeks Intermediate Code Generation")
    add_para("6. Oracle Java Docs")
    add_para("7. IEEE CFG Analysis Paper")

    doc.save("Final_Report_Case_Study_2_V3.docx")
    print("Final DOCX updated with extended content and CFG diagram successfully!")

if __name__ == "__main__":
    main()
